from docx import Document
from docx.enum.section import WD_ORIENTATION
from database import connect
from algorithmen import nums


class Doc:
    def __init__(self, date: str):
        self.doc = Document()
        self.date = date
        self.section = self.doc.sections[0]
        self.width = self.section.page_width
        self.height = self.section.page_height

    def create(self):
        self.section.orientation = WD_ORIENTATION.LANDSCAPE
        self.section.page_width = self.height
        self.section.page_height = self.width

        # set date
        date = self.doc.add_paragraph(self.date)
        date.alignment = 2

        # add headline
        self.doc.add_heading("Fischereischeinprüfung", 1)

        # connect to database
        connection = connect.Database().connect()
        cursor = connection.cursor()

        table = self.doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Frage'
        hdr_cells[1].text = 'Antwort 1'
        hdr_cells[2].text = 'Antwort 2'
        hdr_cells[3].text = 'Antwort 3'

        counter = 1
        while counter < 6:
            numbers = nums.Nums().generate()

            while len(numbers) > 0:
                # data = cursor.execute(f"SELECT Question, AnswerA, AnswerB, AnswerC FROM Brandenburg WHERE NR={numbers[-1]} AND SG={counter}")
                cursor.execute(
                    f"SELECT Question, AnswerA, AnswerB, AnswerC FROM Brandenburg WHERE NR=1 AND SG=1")

                data = cursor.fetchone()
                numbers.pop()

                row_cells = table.add_row().cells
                row_cells[0].text = str(data[0])
                row_cells[1].text = str(data[1])
                row_cells[2].text = str(data[2])
                row_cells[3].text = str(data[3])

            counter += 1

        # self.doc.save(f"Prüfung-{self.date}.docx")
